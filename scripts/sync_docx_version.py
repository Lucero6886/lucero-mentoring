#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Đồng bộ nhãn phiên bản in trong các tài liệu .docx với file VERSION.

Vì sao cần: sáu tài liệu chính sách là bản gốc soạn trong Word, và dòng đầu mỗi
file có ghi tay *"… | Version X.Y.Z | DD/MM/YYYY"*. Word không biết gì về file
`VERSION` ở gốc repo, nên mỗi lần nâng phiên bản là sáu nhãn đó bị bỏ quên —
đúng lỗi đã xảy ra một lần ở v1.4.0 và lặp lại ở v1.9.0.

Script này đóng hẳn lớp lỗi đó: nó đọc `VERSION`, đọc ngày phát hành từ
`06_Data/project_portfolio.json` (`meta.updated`), rồi viết lại đúng cụm
`Version …` và ngày trong .docx. Không đụng tới bất kỳ nội dung nào khác.

Chạy TRƯỚC `export_governance_md.py`, vì bản .md được sinh ra từ .docx:

    python3 scripts/sync_docx_version.py
    python3 scripts/export_governance_md.py

Yêu cầu: python-docx (`pip install python-docx`).
"""
import json, pathlib, re, sys

BASE = pathlib.Path(__file__).resolve().parents[1]

DOCS = [
    "01_Governance/Master_Mentoring_Handbook.docx",
    "01_Governance/Mentor_Student_Working_Agreement.docx",
    "01_Governance/AI_and_Academic_Integrity_Policy.docx",
    "01_Governance/Cohort_HK1_2026_2027_Implementation_Guide.docx",
    "03_Operations/Mentoring_Operating_Procedure_SOP.docx",
    "03_Operations/Weekly_Report_and_Meeting_Template.docx",
]

VER_RE = re.compile(r"(Version\s+)\d+\.\d+(?:\.\d+)?")
DATE_RE = re.compile(r"\b\d{2}/\d{2}/\d{4}\b")


def dmy(iso):
    y, m, d = str(iso).split("-")
    return f"{d}/{m}/{y}"


def fix_runs(paragraphs, ver, date):
    """Sửa tại chỗ từng run có chứa nhãn phiên bản. Trả về số run đã đổi."""
    n = 0
    for p in paragraphs:
        if "Version" not in p.text:
            continue
        for r in p.runs:
            if not VER_RE.search(r.text):
                continue
            new = VER_RE.sub(rf"\g<1>{ver}", r.text)
            new = DATE_RE.sub(date, new)
            if new != r.text:
                r.text = new
                n += 1
    return n


def main():
    try:
        import docx
    except ImportError:
        print("LỖI: thiếu python-docx. Cài bằng: pip install python-docx")
        return 1

    vf = BASE / "VERSION"
    if not vf.is_file():
        print("LỖI: thiếu file VERSION ở gốc repo.")
        return 1
    ver = vf.read_text(encoding="utf-8").strip()
    meta = json.loads((BASE / "06_Data" / "project_portfolio.json").read_text(encoding="utf-8"))["meta"]
    date = dmy(meta["updated"])

    changed, missing, untouched = 0, [], []
    for rel in DOCS:
        src = BASE / rel
        if not src.is_file():
            missing.append(rel)
            continue
        d = docx.Document(str(src))
        n = fix_runs(d.paragraphs, ver, date)
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    n += fix_runs(c.paragraphs, ver, date)
        for s in d.sections:
            n += fix_runs(s.header.paragraphs, ver, date)
            n += fix_runs(s.footer.paragraphs, ver, date)
        if n:
            d.save(str(src))
            print(f"  đồng bộ: {rel}  ({n} chỗ)")
            changed += 1
        else:
            untouched.append(rel)

    for m in untouched:
        print(f"  đã đúng sẵn: {m}")
    for m in missing:
        print(f"  ✗ thiếu file: {m}")
    print(f"XONG — nhãn phiên bản {ver} · ngày {date} · sửa {changed}/{len(DOCS)} tài liệu.")
    if changed:
        print("      → chạy tiếp: python3 scripts/export_governance_md.py")
    return 1 if missing else 0


if __name__ == "__main__":
    sys.exit(main())
